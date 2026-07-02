from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_ROOT = Path(
    r"C:\Users\tanli\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))

from table_geometry import apply_table_geometry  # noqa: E402


OUTPUT_PATH = Path(r"D:\user_wuyou\lilishop-master\lilishop-master\docs\批发商城APP联调接口文档.docx")
ARIAL = "Arial"


def set_run_font(run, *, size=11, bold=None, italic=None, color="000000"):
    run.font.name = ARIAL
    run._element.rPr.rFonts.set(qn("w:ascii"), ARIAL)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ARIAL)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, *, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, bold=False, size=10.5, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    set_spacing(paragraph, before=0, after=0, line=1.12)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc, text, *, size=11, bold=False, italic=False, color="000000", before=0, after=6):
    paragraph = doc.add_paragraph()
    set_spacing(paragraph, before=before, after=after, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(paragraph, before=0, after=3, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=26, color="000000")


def add_h1(doc, text):
    add_paragraph(doc, text, size=20, before=20, after=6)


def add_h2(doc, text):
    add_paragraph(doc, text, size=16, before=18, after=6)


def add_h3(doc, text):
    add_paragraph(doc, text, size=14, color="434343", before=16, after=4)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_spacing(paragraph, before=0, after=4, line=1.15)
        if isinstance(item, tuple):
            lead, detail = item
            run = paragraph.add_run(f"{lead}：")
            set_run_font(run, size=11, bold=True)
            run = paragraph.add_run(detail)
            set_run_font(run, size=11)
        else:
            run = paragraph.add_run(item)
            set_run_font(run, size=11)


def add_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for ridx, row in enumerate(rows):
        for cidx, value in enumerate(row):
            cell = table.rows[ridx].cells[cidx]
            is_header = ridx == 0
            set_cell_text(cell, value, bold=is_header, size=10 if is_header else 10.5)
            if is_header:
                shade_cell(cell, "F3F4F6")
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    doc.add_paragraph()
    return table


def endpoint_block(doc, title, method, path, auth, data_source, purpose, request_rows, response_rows, notes):
    add_h3(doc, title)
    add_table(
        doc,
        [
            ["项", "说明"],
            ["请求方式", method],
            ["接口路径", path],
            ["是否登录", auth],
            ["数据来源", data_source],
            ["接口用途", purpose],
        ],
        [1700, 7660],
    )
    if request_rows:
        add_paragraph(doc, "请求参数说明")
        add_table(doc, [["字段", "位置/类型", "必填", "说明"]] + request_rows, [1800, 1700, 900, 4960])
    if response_rows:
        add_paragraph(doc, "返回字段说明")
        add_table(doc, [["字段", "类型", "说明"]] + response_rows, [3600, 1400, 4360])
    if notes:
        add_paragraph(doc, "联调注意")
        add_bullets(doc, notes)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = ARIAL
    normal._element.rPr.rFonts.set(qn("w:ascii"), ARIAL)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), ARIAL)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)
    normal.font.size = Pt(11)

    add_title(doc, "批发商城 APP 联调接口文档")
    add_paragraph(doc, "适用对象：APP 前端 / H5 前端 / 联调测试", size=10.5, color="555555", after=8)
    add_paragraph(
        doc,
        "整理原则：按 UI 页面分目录；字段说明写清含义、状态和传值；明确区分 ES、缓存、数据库接口，避免前端误连慢接口或误用旧接口。",
        size=10.5,
        color="555555",
        after=8,
    )

    add_h1(doc, "1. 联调总说明")
    add_bullets(
        doc,
        [
            ("统一返回", "所有接口统一使用 ResultMessage<T>。success 表示业务成功与否；code 为业务码；message 为提示文案；result 才是真实业务数据。前端不要只看 HTTP 200，要同时看 success/code。"),
            ("登录令牌", "业务接口优先使用 Authorization: Bearer <accessToken>。登录流程本身先拿 loginSessionToken，再选身份换取最终 accessToken/refreshToken。"),
            ("分页参数", "列表接口统一使用 pageNumber、pageSize、sort、order、notConvert。pageNumber 从 1 开始，pageSize 默认 10，项目里最大 pageSize 为 100。"),
            ("客户端类型", "首页装修和版本接口常用 clientType，APP 联调建议传 APP。不要把 PC/H5 的返回结构直接复用到 APP 页面。"),
        ],
    )

    add_h2(doc, "1.1 一定要先看：ES、缓存、数据库的使用边界")
    add_bullets(
        doc,
        [
            ("商品搜索页必须走 ES 接口", "商品列表搜索、筛选、品牌/分类/属性联动，要优先使用 /buyer/goods/goods/es 和 /buyer/goods/goods/es/related。这一套已经是为性能优化做好的搜索链路。不要用 /buyer/goods/goods 代替搜索页。"),
            ("商品详情存在缓存", "商品详情 GoodsVO 走 CachePrefix.GOODS，缓存时长 7200 秒；SKU 详情也有缓存和活动价合并逻辑。前端出现“刚改后台价格但详情页未立刻变更”时，先确认缓存刷新和上架索引刷新，不要误判接口错。"),
            ("购物车与结算上下文是缓存态", "购物车、地址选择、配送方式、优惠券选择、发票选择和 createTrade 之间不是纯数据库直读，而是基于购物车缓存上下文逐步渲染。前端必须按接口顺序调用，不要本地拼结算金额。"),
            ("热词走 Redis", "搜索热词 /buyer/goods/goods/hot-words 是 Redis 热词集合，不是数据库表直接列表。"),
            ("首页有两套接口", "结构化首页优先用 /buyer/other/home/platform；如果是读取装修 JSON 再本地渲染，可用 /buyer/other/pageData/getIndex。两套返回结构不同，前端不要混用字段。"),
            ("登录会话走 Redis", "短信认证/一键登录成功后返回的 loginSessionToken 默认只保留 10 分钟，存在 Redis 登录会话里。身份页停留太久导致选择身份失败，属于会话过期，不是前端 token 丢失。"),
        ],
    )

    add_h2(doc, "1.2 响应和分页字段")
    add_table(
        doc,
        [
            ["字段", "类型", "说明"],
            ["success", "boolean", "业务是否成功；true 成功，false 失败。"],
            ["code", "int", "统一业务码；200 通常表示成功，其他值表示具体失败原因。"],
            ["message", "string", "提示文案；失败时优先结合 code 判断，不要只依赖中文文案。"],
            ["timestamp", "long", "服务端响应时间戳，毫秒。"],
            ["result", "object", "真实业务数据载荷。"],
            ["pageNumber", "int", "分页页码，从 1 开始。"],
            ["pageSize", "int", "分页大小，默认 10，建议不要超过 100。"],
            ["sort", "string", "排序字段，默认按驼峰转下划线。"],
            ["order", "string", "排序方向，asc 或 desc。"],
            ["notConvert", "boolean", "是否关闭 sort 的驼峰转下划线；ES 查询常会设为 true。"],
        ],
        [1700, 1200, 6460],
    )

    add_h1(doc, "2. 登录与角色选择")
    endpoint_block(
        doc,
        "2.1 短信认证登录",
        "POST",
        "/common/auth/login/sms",
        "否",
        "Redis 登录会话 + Member 账号自动创建",
        "校验短信验证码后创建登录会话，不直接签发业务 token。",
        [
            ["uuid", "Header / string", "是", "短信验证码会话标识，必须和发送验证码时的 uuid 对应。"],
            ["mobile", "Body / string", "是", "手机号；如果系统中还没有这个手机号，会自动注册会员。"],
            ["code", "Body / string", "是", "短信验证码。"],
        ],
        [
            ["loginSessionToken", "string", "登录会话 token，后续查身份列表和选身份都靠它；默认 10 分钟过期。"],
            ["memberId", "string", "当前手机号对应的会员 ID。"],
            ["mobile", "string", "本次认证确认的手机号。"],
            ["channel", "enum", "认证渠道：SMS 或 ONE_CLICK。"],
            ["identityOptions", "array", "当前账号可进入的身份列表。"],
        ],
        [
            ("不是最终登录", "这个接口返回的不是 accessToken，前端不能拿 loginSessionToken 直接访问买家/商家业务接口。"),
            ("过期处理", "如果身份选择页停留过久，loginSessionToken 过期要重新做短信认证。"),
        ],
    )
    endpoint_block(
        doc,
        "2.2 本机号码一键登录认证",
        "POST",
        "/common/auth/login/one-click",
        "否",
        "运营商凭证解析 + Redis 登录会话",
        "用运营商一键登录凭证换手机号并创建登录会话。",
        [["loginToken", "Body / string", "是", "运营商或服务商返回的一键登录临时凭证。"]],
        [
            ["loginSessionToken", "string", "后续选身份要用的登录会话 token。"],
            ["identityOptions", "array", "当前手机号可进入的身份列表。"],
        ],
        [("失败理解", "一键登录失败通常是凭证无效、已过期或短信配置不支持，不是账号异常。")],
    )
    endpoint_block(
        doc,
        "2.3 查询登录身份列表",
        "GET",
        "/common/auth/login/identities/{loginSessionToken}",
        "否",
        "Redis 登录会话 + 身份判定服务",
        "获取当前手机号可进入的消费者、代理商、供货商身份状态。",
        [["loginSessionToken", "Path / string", "是", "短信或一键登录成功后返回的登录会话 token。"]],
        [
            ["identityCode", "enum", "身份编码：CONSUMER 消费者、AGENT 代理商、SUPPLIER 供货商。"],
            ["authDomain", "enum", "认证域：MEMBER 买家域，STORE 供货商域。"],
            ["available", "boolean", "是否允许当前会话直接进入该身份。"],
            ["status", "enum", "AVAILABLE 可进入；NOT_OPENED 未开通；PENDING 审核中；REJECTED 驳回；DISABLED 禁用。"],
            ["message", "string", "给前端展示的状态文案。"],
            ["nextAction", "enum", "NONE / APPLY_AGENT / APPLY_STORE / RESUBMIT / CONTACT_ADMIN。"],
        ],
        [
            ("代理商与供货商不是一定都有", "消费者身份通常可直接进入；代理商和供货商要看角色关系、店铺类型、审核状态、启停状态。"),
            ("前端展示建议", "当 available=false 时，不要直接报错，可按 nextAction 做“去申请 / 重新提交 / 联系管理员”引导。"),
        ],
    )
    endpoint_block(
        doc,
        "2.4 买家端确认身份并换取 token",
        "POST",
        "/buyer/passport/member/selectIdentity",
        "否",
        "登录会话 + 会员域 token 签发",
        "消费者端和代理商端最终登录入口。",
        [
            ["loginSessionToken", "Body / string", "是", "登录会话 token。"],
            ["identityCode", "Body / enum", "是", "可传 CONSUMER 或 AGENT。"],
        ],
        [
            ["identityCode", "enum", "本次进入的身份。"],
            ["authDomain", "enum", "MEMBER，表示买家域 token。"],
            ["token", "object", "最终 accessToken / refreshToken。"],
            ["memberProfile", "object", "当前会员基础资料。"],
            ["identityStatusSummary", "object", "重新计算后的全量身份状态快照。"],
        ],
        [("Bearer token 从这里开始生效", "之后访问买家接口需要带 Authorization: Bearer <accessToken>。")],
    )
    endpoint_block(
        doc,
        "2.5 商家端确认身份并换取 token",
        "POST",
        "/store/passport/login/selectIdentity",
        "否",
        "登录会话 + 商家域 token 签发",
        "供货商端最终登录入口。",
        [
            ["loginSessionToken", "Body / string", "是", "登录会话 token。"],
            ["identityCode", "Body / enum", "是", "固定传 SUPPLIER。"],
        ],
        [
            ["authDomain", "enum", "STORE，表示商家域 token。"],
            ["token", "object", "最终访问商家接口的 token。"],
        ],
        [("接口域不要混", "SUPPLIER 身份拿到的是 STORE 域 token，不能直接访问 buyer 域仅会员接口。")],
    )

    add_h1(doc, "3. 首页与导航页面")
    endpoint_block(
        doc,
        "3.1 结构化首页",
        "GET",
        "/buyer/other/home/platform",
        "否",
        "DB + 运营配置 + 推荐聚合 + 缓存友好结构化输出",
        "APP 首页主接口，适合直接按字段渲染轮播、金刚区、楼层、秒杀、热销、新品、常买店铺。",
        [["clientType", "Query / string", "是", "客户端类型，APP 联调建议固定传 APP。"]],
        [
            ["banners", "array", "首页轮播图列表。"],
            ["shortcutNavList", "array", "金刚区快捷入口。"],
            ["floorModules", "array", "楼层模块。"],
            ["seckill", "object", "秒杀模块。"],
            ["monthlyHotGoods", "array", "月度热销商品卡片。"],
            ["newGoods", "array", "上新商品卡片。"],
            ["frequentStores", "array", "常买店铺。"],
        ],
        [
            ("优先使用这个接口", "如果首页是按固定结构渲染，优先用这个接口，而不是自己解析装修 JSON。"),
            ("常买店铺不是运营手填", "它是按订单、关注、足迹等规则综合生成，和后台静态配置不是一回事。"),
        ],
    )
    endpoint_block(
        doc,
        "3.2 首页装修原始数据",
        "GET",
        "/buyer/other/pageData/getIndex",
        "否",
        "DB 装修页数据",
        "返回原始装修 pageData JSON，适合页面引擎自己解释渲染。",
        [["clientType", "Query / string", "是", "客户端类型。"]],
        [["pageData", "string(JSON)", "首页装修 JSON 字符串。"]],
        [("和结构化首页二选一", "前端不要把 /platform 和 /getIndex 的字段混着取。")],
    )
    endpoint_block(
        doc,
        "3.3 分类树",
        "GET",
        "/buyer/goods/category/get/{parentId}",
        "否",
        "DB 分类表",
        "分类页和三级分类下拉的基础接口。",
        [["parentId", "Path / string", "是", "父级分类 ID；根节点按项目约定传顶级 ID。"]],
        [
            ["id", "string", "分类 ID。"],
            ["name", "string", "分类名称。"],
            ["level", "int", "层级。"],
            ["children", "array", "下级分类集合，具体是否返回取决于实现。"],
        ],
        [("三级分类说明", "设计图说明分类类似京东三级分类，前端要支持分级加载或三级展开。")],
    )
    endpoint_block(
        doc,
        "3.4 搜索热词",
        "GET",
        "/buyer/goods/goods/hot-words",
        "否",
        "Redis ZSet 热词缓存",
        "搜索页默认展示热词。",
        [["count", "Query / int", "否", "返回热词数量，建议前端传 10 或 20。"]],
        [["result", "array<string>", "热词列表。"]],
        [("不是数据库字典", "热词是动态排行，刷新频率可能比普通配置更高。")],
    )

    add_h1(doc, "4. 商品搜索、商品详情、店铺详情")
    endpoint_block(
        doc,
        "4.1 商品搜索列表（ES）",
        "GET",
        "/buyer/goods/goods/es",
        "否",
        "Elasticsearch 索引 + 搜索降级逻辑",
        "商品列表搜索、关键词搜索、筛选排序主接口。",
        [
            ["keyword", "Query / string", "否", "搜索关键字。"],
            ["categoryId", "Query / string", "否", "分类筛选。"],
            ["brandId", "Query / string", "否", "品牌筛选。"],
            ["storeId", "Query / string", "否", "店铺筛选。"],
            ["pageNumber", "Query / int", "否", "页码，默认 1。"],
            ["pageSize", "Query / int", "否", "每页条数，默认 10。"],
            ["sort", "Query / string", "否", "排序字段，ES 查询常配合 notConvert=true。"],
            ["order", "Query / string", "否", "asc 或 desc。"],
        ],
        [
            ["records", "array", "商品记录列表。"],
            ["goodsId", "string", "商品 ID。"],
            ["id", "string", "ES 中常用 SKU/索引主键，前端联调时注意与 goodsId 区分。"],
            ["goodsName", "string", "商品名称。"],
            ["thumbnail", "string", "商品主图。"],
            ["price", "number", "当前展示售价。"],
            ["promotionPrice", "number", "促销价；是否展示要结合 promotionFlag。"],
            ["storeName", "string", "店铺名称。"],
            ["salesModel", "enum", "RETAIL 零售 / WHOLESALE 批发。"],
        ],
        [
            ("搜索页必须用它", "这是已经为性能优化好的 ES 接口；如果改回 DB 列表，会把搜索和筛选压力重新打回数据库。"),
            ("排序字段别乱传", "项目里 PageVO 默认会做驼峰转下划线；ES 查询一般会显式 notConvert=true，所以前端只传接口文档约定字段。"),
        ],
    )
    endpoint_block(
        doc,
        "4.2 搜索筛选项（ES 相关聚合）",
        "GET",
        "/buyer/goods/goods/es/related",
        "否",
        "Elasticsearch 聚合",
        "返回当前搜索结果下可联动的分类、品牌、属性筛选项。",
        [
            ["keyword", "Query / string", "否", "和搜索页保持同一组条件。"],
            ["categoryId", "Query / string", "否", "当前已选分类。"],
        ],
        [
            ["categoryList", "array", "相关分类聚合。"],
            ["brandList", "array", "相关品牌聚合。"],
            ["attrList / paramList", "array", "相关属性或参数聚合。"],
        ],
        [("必须跟搜索条件同传", "否则筛选面板和结果列表会不一致。")],
    )
    endpoint_block(
        doc,
        "4.3 商品列表（DB）",
        "GET",
        "/buyer/goods/goods",
        "否",
        "数据库查询",
        "普通商品分页列表，适合后台弱搜索场景或部分非搜索页面。",
        [
            ["goodsName / keyword", "Query / string", "否", "商品关键词。"],
            ["categoryId", "Query / string", "否", "分类。"],
        ],
        [["records", "array", "Goods 实体分页记录。"]],
        [("不要拿来做主搜索页", "它不是 ES 搜索链路，筛选和性能都不如 ES 版本。")],
    )
    endpoint_block(
        doc,
        "4.4 商品详情",
        "GET",
        "/buyer/goods/goods/get/{goodsId}",
        "否",
        "DB + Redis 商品缓存",
        "商品详情主接口，返回 GoodsVO。",
        [["goodsId", "Path / string", "是", "商品 ID。"]],
        [
            ["goodsId", "string", "商品 ID。"],
            ["goodsName", "string", "商品名称。"],
            ["thumbnail / goodsGalleryList", "string/array", "主图与轮播图。"],
            ["storeId", "string", "所属店铺 ID。"],
            ["storeName", "string", "所属店铺名。"],
            ["price", "number", "基础售价。"],
            ["salesModel", "enum", "销售模式：RETAIL / WHOLESALE。"],
            ["goodsParamsDTOList", "array", "商品参数组。"],
            ["mobileIntro / intro", "string", "图文详情。"],
        ],
        [("有缓存", "商品详情有 7200 秒缓存；后台改完商品后如前端短时间没变，先排查缓存和索引刷新。")],
    )
    endpoint_block(
        doc,
        "4.5 商品详情 SKU 维度信息",
        "GET",
        "/buyer/goods/goods/sku/{goodsId}/{skuId}",
        "否",
        "SKU 缓存 + 活动价/促销合并 + DB 回源",
        "规格切换时获取具体 SKU 的价格、库存、活动信息。",
        [
            ["goodsId", "Path / string", "是", "商品 ID。"],
            ["skuId", "Path / string", "是", "SKU ID。"],
        ],
        [
            ["goodsSku", "object", "当前 SKU 基础信息。"],
            ["price", "number", "当前购买价。"],
            ["quantity", "int", "当前库存。"],
            ["promotionMap", "object", "当前命中的促销活动集合。"],
            ["wholesaleList", "array", "批发阶梯价规则。"],
        ],
        [("价格以这个接口为准", "详情页切换规格后，前端不要继续沿用 goods 级默认价格。")],
    )
    endpoint_block(
        doc,
        "4.6 店铺详情",
        "GET",
        "/buyer/store/store/get/{storeId}",
        "否",
        "DB + 店铺缓存（如命中）",
        "店铺详情页基础信息。",
        [["storeId", "Path / string", "是", "店铺 ID。"]],
        [
            ["storeId", "string", "店铺 ID。"],
            ["storeName", "string", "店铺名。"],
            ["storeLogo", "string", "店铺 logo。"],
            ["storeDesc", "string", "店铺简介。"],
            ["favoriteNum", "int", "关注量。"],
        ],
        [("和商品搜索 storeId 打通", "前端从搜索页、商品详情页跳转店铺详情时都要保留 storeId。")],
    )

    add_h1(doc, "5. 购物车、结算、订单、售后")
    endpoint_block(
        doc,
        "5.1 加入购物车",
        "POST",
        "/buyer/trade/carts",
        "是",
        "购物车缓存上下文",
        "向购物车加入一个 SKU。",
        [
            ["skuId", "Query/Form / string", "是", "SKU ID。"],
            ["num", "Query/Form / int", "是", "购买数量，必须大于 0。"],
            ["cartType", "Query/Form / string", "否", "购物车类型，常见值：CART / BUY_NOW / PINTUAN / POINTS。"],
        ],
        [],
        [("cartType 很关键", "普通购物车、立即购买、拼团购买、积分购买是不同上下文，前端不能混传。")],
    )
    endpoint_block(
        doc,
        "5.2 获取购物车详情",
        "GET",
        "/buyer/trade/carts/all",
        "是",
        "购物车缓存渲染结果",
        "购物车页面主接口。",
        [],
        [
            ["cartList", "array", "按店铺分组的购物车。"],
            ["skuList", "array", "店铺下 SKU 列表。"],
            ["checked", "boolean", "是否选中。"],
            ["invalid", "boolean", "是否失效。"],
            ["errorMessage", "string", "失效原因，如已下架、库存不足。"],
            ["priceDetailDTO", "object", "价格明细。"],
        ],
        [("金额不要前端自己算", "优惠、运费、活动价都在服务端渲染链路里。")],
    )
    endpoint_block(
        doc,
        "5.3 获取结算页详情",
        "GET",
        "/buyer/trade/carts/checked",
        "是",
        "购物车缓存渲染结果",
        "结算页主接口。",
        [["way", "Query / string", "是", "CART / BUY_NOW / PINTUAN / POINT 等结算方式。"]],
        [
            ["cartList", "array", "结算店铺分组。"],
            ["memberAddress", "object", "当前选择的收货地址。"],
            ["storeAddress", "object", "当前选择的自提地址。"],
            ["canReceiveCoupon", "array", "当前可领取或可用优惠券。"],
            ["priceDetailVO", "object", "结算价格汇总。"],
        ],
        [("结算顺序要对", "通常先拿 checked，再选地址、配送方式、优惠券、发票，最后 createTrade。")],
    )
    endpoint_block(
        doc,
        "5.4 选择收货地址",
        "GET",
        "/buyer/trade/carts/shippingAddress",
        "是",
        "购物车缓存上下文",
        "为当前结算方式绑定收货地址。",
        [
            ["shippingAddressId", "Query / string", "是", "收货地址 ID。"],
            ["way", "Query / string", "是", "结算方式。"],
        ],
        [],
        [("不是全局设置", "它是写入当前购物车/交易上下文，不是改用户默认地址。")],
    )
    endpoint_block(
        doc,
        "5.5 选择配送方式",
        "PUT",
        "/buyer/trade/carts/shippingMethod",
        "是",
        "购物车缓存上下文",
        "切换物流、自提、同城等配送方式。",
        [
            ["shippingMethod", "Query/Form / string", "是", "SELF_PICK_UP / LOCAL_TOWN_DELIVERY / LOGISTICS。"],
            ["way", "Query/Form / string", "是", "结算方式。"],
        ],
        [],
        [("配送方式影响价格", "切换后要重新拉结算信息，不能只本地改文案。")],
    )
    endpoint_block(
        doc,
        "5.6 选择优惠券",
        "GET",
        "/buyer/trade/carts/select/coupon",
        "是",
        "购物车缓存上下文 + 优惠券缓存/规则",
        "在当前结算上下文中选中或取消优惠券。",
        [
            ["way", "Query / string", "是", "结算方式。"],
            ["memberCouponId", "Query / string", "是", "会员券 ID。"],
            ["used", "Query / boolean", "是", "true 使用；false 取消使用。"],
        ],
        [],
        [("传会员券 ID 不是活动券 ID", "要区分 couponId 和 memberCouponId。")],
    )
    endpoint_block(
        doc,
        "5.7 创建交易",
        "POST",
        "/buyer/trade/carts/create/trade",
        "是",
        "购物车缓存上下文 -> 订单创建",
        "真正提交订单前的交易创建接口。",
        [["tradeParams", "Body / object", "是", "交易参数对象；常见包含 way、client、parentOrderSn、备注等。"]],
        [
            ["orderSn / orderList", "string/array", "创建成功后的订单信息。"],
            ["payPrice / flowPrice", "number", "应付金额。"],
        ],
        [("重复提交要防抖", "接口有防重复提交保护，前端仍建议按钮 loading+禁点。")],
    )
    endpoint_block(
        doc,
        "5.8 地址列表",
        "GET",
        "/buyer/member/address",
        "是",
        "DB 地址表",
        "地址管理页面列表。",
        [
            ["pageNumber", "Query / int", "否", "页码。"],
            ["pageSize", "Query / int", "否", "分页大小。"],
        ],
        [
            ["id", "string", "地址 ID。"],
            ["name / consignee", "string", "收货人。"],
            ["mobile", "string", "手机号。"],
            ["consigneeAddressPath", "string", "地区名称路径，常为省/市/区/街道拼接。"],
            ["consigneeAddressIdPath", "string", "地区 ID 路径。"],
            ["detail", "string", "详细地址。"],
            ["isDefault", "boolean", "是否默认地址。"],
        ],
        [("地址模型不是省市区分字段优先", "项目里常用 path/pathId 的方式保存地区路径，前端不要假设只有 provinceId/cityId/areaId。")],
    )
    endpoint_block(
        doc,
        "5.9 新增地址",
        "POST",
        "/buyer/member/address",
        "是",
        "DB 地址表",
        "新增用户收货地址。",
        [
            ["name / consignee", "Body / string", "是", "收货人姓名。"],
            ["mobile", "Body / string", "是", "手机号。"],
            ["consigneeAddressPath", "Body / string", "是", "地区名称路径。"],
            ["consigneeAddressIdPath", "Body / string", "是", "地区 ID 路径。"],
            ["detail", "Body / string", "是", "详细地址。"],
            ["isDefault", "Body / boolean", "否", "是否默认地址；不传时后端默认 false。"],
        ],
        [["id", "string", "新建成功后的地址 ID。"]],
        [("前端一定传完整路径", "仅传最后一级 areaId 往往不够，建议把地区名称路径和 ID 路径都带上。")],
    )
    endpoint_block(
        doc,
        "5.10 订单列表",
        "GET",
        "/buyer/order/order",
        "是",
        "DB 订单查询",
        "我的订单页主接口，消费者和代理商都使用买家端订单接口查看本人订单。",
        [
            ["orderStatus / tag", "Query / string", "否", "按状态筛选，如 WAIT_PAY、WAIT_SHIP、WAIT_ROG、COMPLETE、CANCELLED。"],
            ["keyword", "Query / string", "否", "订单号或商品关键词。"],
            ["pageNumber", "Query / int", "否", "分页页码。"],
            ["pageSize", "Query / int", "否", "分页大小。"],
        ],
        [
            ["orderSn", "string", "订单号。"],
            ["orderStatus", "enum", "UNPAID / PAID / DELIVERED / COMPLETED / CANCELLED 等。"],
            ["orderStatusValue / allowOperation", "object", "状态文案和可操作集合。"],
            ["flowPrice", "number", "订单金额。"],
            ["groupNum / groupBuy", "number/object", "拼团类页面常用字段。"],
        ],
        [("状态文案以后端为准", "不要前端自行重新命名已付款待发货、已发货待收货等状态。")],
    )
    endpoint_block(
        doc,
        "5.11 订单详情",
        "GET",
        "/buyer/order/order/{orderSn}",
        "是",
        "DB 订单详情 + 包裹/日志聚合",
        "订单详情页面主接口。",
        [["orderSn", "Path / string", "是", "订单号。"]],
        [
            ["order", "object", "订单主信息。"],
            ["orderItems", "array", "订单商品列表。"],
            ["priceDetailVO", "object", "金额拆分明细。"],
            ["logList", "array", "操作日志。"],
            ["allowOperation", "object", "当前订单允许的动作。"],
        ],
        [("详情页金额以这里为准", "设计图里的商品金额、优惠、运费、实付金额都要从详情接口取，不要从列表页拼。")],
    )
    endpoint_block(
        doc,
        "5.12 确认收货",
        "POST",
        "/buyer/order/order/{orderSn}/receiving",
        "是",
        "DB 订单状态流转",
        "待收货订单确认收货。",
        [["orderSn", "Path / string", "是", "订单号。"]],
        [],
        [("只有 DELIVERED 才能收货", "如果订单不在待收货状态，后端会拒绝。")],
    )
    endpoint_block(
        doc,
        "5.13 取消订单",
        "POST",
        "/buyer/order/order/{orderSn}/cancel",
        "是",
        "DB 订单状态流转",
        "取消未完成订单。",
        [
            ["orderSn", "Path / string", "是", "订单号。"],
            ["reason", "Query / string", "是", "取消原因，建议前端提供原因选项和自定义补充。"],
        ],
        [],
        [("原因必传", "不要只调接口不带 reason。")],
    )
    endpoint_block(
        doc,
        "5.14 售后列表",
        "GET",
        "/buyer/order/afterSale/page",
        "是",
        "DB 售后单查询",
        "售后列表页面主接口。",
        [
            ["pageNumber", "Query / int", "否", "页码。"],
            ["pageSize", "Query / int", "否", "每页条数。"],
            ["serviceType", "Query / string", "否", "售后类型筛选。"],
            ["status", "Query / string", "否", "售后状态筛选。"],
        ],
        [
            ["sn", "string", "售后单号。"],
            ["serviceType", "enum", "RETURN_MONEY 仅退款 / RETURN_GOODS 退货退款。"],
            ["status", "string", "售后状态。"],
            ["applyRefundPrice", "number", "申请金额。"],
        ],
        [("售后类型目前只有两个", "实际代码枚举只有 RETURN_MONEY 和 RETURN_GOODS，前端不要自造“换货”状态。")],
    )
    endpoint_block(
        doc,
        "5.15 申请售后页信息",
        "GET",
        "/buyer/order/afterSale/applyAfterSaleInfo/{sn}",
        "是",
        "DB 订单项 + 售后资格校验",
        "进入售后申请页前先查可售后信息。",
        [["sn", "Path / string", "是", "订单货物编号，不是订单号。"]],
        [
            ["goodsInfo", "object", "当前订单项商品信息。"],
            ["maxRefundAmount", "number", "最大可申请退款金额。"],
            ["serviceTypeList", "array", "当前可申请售后类型。"],
        ],
        [("这里的 sn 是订单货物编号", "不是 orderSn，前端容易传错。")],
    )
    endpoint_block(
        doc,
        "5.16 提交售后申请",
        "POST",
        "/buyer/order/afterSale/save/{orderItemSn}",
        "是",
        "DB 售后创建",
        "提交售后申请。",
        [
            ["orderItemSn", "Path / string", "是", "订单货物编号；虽然路径中带这个参数，但控制器实际接收 DTO 表单字段，前端请按接口调试结果一起带。"],
            ["serviceType", "Body/Form / string", "是", "RETURN_MONEY 或 RETURN_GOODS。"],
            ["reason", "Body/Form / string", "是", "售后原因。"],
            ["problemDesc", "Body/Form / string", "否", "问题描述。"],
            ["price", "Body/Form / number", "是", "申请金额。"],
            ["images", "Body/Form / array", "否", "凭证图片。"],
        ],
        [
            ["sn", "string", "售后单号。"],
            ["status", "string", "新建后的售后状态。"],
        ],
        [("这是一个历史风格接口", "路径带 orderItemSn，但方法签名主要吃 DTO，联调时建议按 swagger/抓包结果把路径和表单一起带全。")],
    )

    add_h1(doc, "6. 拼团、秒杀、优惠券、积分、收藏、足迹、消息、钱包")
    endpoint_block(
        doc,
        "6.1 拼团列表",
        "GET",
        "/buyer/promotion/pintuan",
        "否",
        "DB + 活动规则",
        "精品好团页面列表。",
        [
            ["pageNumber", "Query / int", "否", "页码。"],
            ["pageSize", "Query / int", "否", "分页大小。"],
        ],
        [
            ["id", "string", "拼团活动 ID。"],
            ["title", "string", "活动标题。"],
            ["goodsId", "string", "商品 ID。"],
            ["pintuanPrice", "number", "拼团价。"],
            ["needMemberNum", "int", "成团人数。"],
            ["promotionStatus", "enum", "NEW / START / END / CLOSE。"],
        ],
        [("设计说明", "原型说明团购类似拼多多，两人起团；前端成团人数展示请以后端返回为准。")],
    )
    endpoint_block(
        doc,
        "6.2 拼团详情",
        "GET",
        "/buyer/promotion/pintuan/{pintuanId}",
        "否",
        "DB + 活动规则 + 商品聚合",
        "拼团详情页。",
        [["pintuanId", "Path / string", "是", "拼团活动 ID。"]],
        [
            ["pintuanPrice", "number", "拼团价。"],
            ["needMemberNum", "int", "成团人数。"],
            ["startTime/endTime", "datetime", "活动时间。"],
            ["goodsInfo", "object", "关联商品信息。"],
        ],
        [("分享页另有接口", "团分享页常配合 /buyer/promotion/pintuan/share 使用。")],
    )
    endpoint_block(
        doc,
        "6.3 秒杀列表",
        "GET",
        "/buyer/promotion/seckill",
        "否",
        "缓存友好活动读取 + SKU 缓存",
        "秒杀列表页。",
        [],
        [
            ["timeline", "string/int", "时间场次。"],
            ["goodsList", "array", "秒杀商品列表。"],
        ],
        [("秒杀价格不要从普通详情页猜", "秒杀活动通常要结合活动接口或 SKU 详情返回判断。")],
    )
    endpoint_block(
        doc,
        "6.4 优惠券列表",
        "GET",
        "/buyer/promotion/coupon",
        "是",
        "DB + 会员券状态缓存/规则",
        "我的优惠券页面。",
        [
            ["pageNumber", "Query / int", "否", "页码。"],
            ["pageSize", "Query / int", "否", "分页大小。"],
        ],
        [
            ["id", "string", "会员券 ID。"],
            ["couponName", "string", "优惠券名称。"],
            ["memberCouponStatus", "enum", "NEW 未使用 / USED 已使用 / EXPIRE 已过期 / CLOSED 已关闭。"],
            ["startTime/endTime", "datetime", "有效期。"],
            ["price / consumeThreshold", "number", "优惠面额与门槛。"],
        ],
        [("前端 tab 建议按状态过滤", "设计图里有可领、未使用、已使用/过期三种视图，要注意可领券和我的券不是同一接口。")],
    )
    endpoint_block(
        doc,
        "6.5 领取优惠券",
        "GET",
        "/buyer/promotion/coupon/receive/{couponId}",
        "是",
        "DB + 领取规则校验",
        "领取活动券。",
        [["couponId", "Path / string", "是", "活动券 ID，不是会员券 ID。"]],
        [],
        [("这是 GET 动作接口", "项目里有一些历史接口虽然是动作，但依然用 GET，前端联调时按真实接口来。")],
    )
    endpoint_block(
        doc,
        "6.6 积分商城列表",
        "GET",
        "/buyer/promotion/pointsGoods",
        "否/部分兑换动作需登录",
        "DB + 积分商品能力",
        "积分商城列表页。",
        [
            ["pageNumber", "Query / int", "否", "页码。"],
            ["pageSize", "Query / int", "否", "分页大小。"],
        ],
        [
            ["id", "string", "积分商品 ID。"],
            ["points", "int", "兑换所需积分。"],
            ["stock", "int", "库存。"],
            ["exchangeLimit", "int", "每人兑换上限。"],
        ],
        [("兑换前确认库存与积分", "前端展示要同时给出积分消耗和库存限制。")],
    )
    endpoint_block(
        doc,
        "6.7 商品收藏",
        "POST",
        "/buyer/member/favorite/goods/{goodsId}",
        "是",
        "DB 收藏表",
        "收藏或取消收藏商品。",
        [["goodsId", "Path / string", "是", "商品 ID。"]],
        [],
        [("通常是切换动作", "具体是收藏还是取消收藏，要结合当前状态和接口实现确认。")],
    )
    endpoint_block(
        doc,
        "6.8 浏览历史",
        "GET",
        "/buyer/member/footprint/history",
        "是",
        "DB 足迹表",
        "浏览历史页面。",
        [],
        [
            ["goodsId", "string", "商品 ID。"],
            ["goodsName", "string", "商品名称。"],
            ["thumbnail", "string", "商品图。"],
            ["visitTime", "datetime", "浏览时间。"],
        ],
        [],
    )
    endpoint_block(
        doc,
        "6.9 钱包首页",
        "GET",
        "/buyer/wallet/wallet",
        "是",
        "DB 钱包表 + 金额汇总",
        "钱包页主接口。",
        [],
        [
            ["memberId", "string", "会员 ID。"],
            ["balance", "number", "钱包余额。"],
            ["canWithdrawAmount / withdrawableAmount", "number", "可提现金额。"],
            ["frozenAmount", "number", "冻结金额。"],
        ],
        [("提现按钮逻辑", "设计图说明消费者只能提现自己充值余额；未录入必要收款信息时，提现按钮应置灰。")],
    )
    endpoint_block(
        doc,
        "6.10 提现申请",
        "POST",
        "/buyer/wallet/withdraw",
        "是",
        "DB 钱包表 + 提现申请单",
        "发起提现。",
        [
            ["withdrawAmount", "Body/Form / number", "是", "提现金额。"],
            ["accountType", "Body/Form / string", "是", "收款账户类型。"],
            ["accountName", "Body/Form / string", "是", "收款人。"],
            ["accountNo", "Body/Form / string", "是", "收款账号。"],
        ],
        [],
        [("金额校验以后端为准", "前端只做基础格式校验，最终可提现额度以后端返回为准。")],
    )
    endpoint_block(
        doc,
        "6.11 站内消息列表（商家消息中心同类）",
        "GET",
        "/store/message/storeMessage 或 /buyer/message/list",
        "是",
        "DB 消息表",
        "消息中心页面。买家端和商家端接口域不同，但联调逻辑相近。",
        [],
        [
            ["id", "string", "消息 ID。"],
            ["title", "string", "消息标题。"],
            ["content", "string", "消息内容。"],
            ["status", "enum", "UN_READY 未读 / ALREADY_READY 已读 / ALREADY_REMOVE 已删除。"],
            ["createTime", "datetime", "消息时间。"],
        ],
        [("状态不要自行映射错", "消息状态枚举和订单状态不是一套，不要复用通用数字码。")],
    )
    endpoint_block(
        doc,
        "6.12 IM 会话列表（跨服务）",
        "GET",
        "/im/talk/list",
        "是",
        "IM 服务数据库 / WebSocket 会话域",
        "聊天会话列表。",
        [],
        [
            ["id / talkId", "string", "会话 ID。"],
            ["lastMessage", "object/string", "最后一条消息。"],
            ["unreadCount", "int", "未读数。"],
        ],
        [("这是独立服务", "聊天接口在 im-api，不在 buyer-api/store-api 里；联调环境和 token 透传要单独确认。")],
    )
    endpoint_block(
        doc,
        "6.13 IM 消息发送（跨服务）",
        "POST",
        "/im/message",
        "是",
        "IM 服务",
        "发送聊天消息。",
        [
            ["talkId", "Body / string", "是", "会话 ID。"],
            ["messageType", "Body / string", "是", "text / image / goods / order 等。"],
            ["content", "Body / string", "是", "消息内容。"],
            ["extraJson", "Body / string", "否", "扩展 JSON。"],
        ],
        [
            ["id", "string", "消息 ID。"],
            ["sendTime", "datetime", "发送时间。"],
        ],
        [("商品卡片和订单卡片建议走 extraJson", "这样前端更容易扩展消息体。")],
    )

    add_h1(doc, "7. 代理商页面")
    endpoint_block(
        doc,
        "7.1 代理商身份检查",
        "GET",
        "/buyer/agent/role/check",
        "是",
        "角色关系表 + 审核状态",
        "进入代理商模块前检查当前会员是否具备代理身份。",
        [],
        [["result", "boolean/object", "是否可进入代理商模块，具体结构以联调返回为准。"]],
        [("和登录身份是两层概念", "登录时能不能选 AGENT 看登录身份服务；进入具体代理业务页仍可能做角色校验。")],
    )
    endpoint_block(
        doc,
        "7.2 我的店铺",
        "GET",
        "/buyer/agent/store",
        "是",
        "代理店铺绑定关系 + 店铺信息",
        "代理商“我的店铺”页面。",
        [],
        [
            ["storeId", "string", "店铺 ID。"],
            ["storeName", "string", "店铺名称。"],
            ["storeLogo", "string", "店铺 logo。"],
            ["bindStatus", "enum", "绑定状态。"],
            ["storeLevel", "string", "店铺等级。"],
            ["auditStatus", "enum", "店铺审核状态。"],
        ],
        [("绑定状态和审核状态要分开显示", "一个表示代理关系，一个表示店铺自身经营状态。")],
    )
    endpoint_block(
        doc,
        "7.3 经营概览",
        "GET",
        "/buyer/agent/dashboard/overview",
        "是",
        "DB 聚合统计",
        "代理商经营概览页。",
        [],
        [
            ["salesAmount", "number", "本月销售额。"],
            ["orderCount", "int", "本月订单数。"],
            ["bindStoreCount", "int", "绑定店铺数。"],
            ["commissionAmount", "number", "累计佣金金额。"],
            ["monthIncome", "number", "本月成交额。"],
            ["todaySales", "int", "今日销量。"],
            ["sevenDaySales", "int", "近 7 日销量。"],
            ["totalOrderCount", "int", "累计订单数。"],
            ["customerUnitPrice", "number", "客单价。"],
            ["returnRate", "number", "退货率，百分比数值。"],
        ],
        [("用于概览卡片", "这些字段适合做顶部 KPI 卡片，不建议前端再按订单列表自己汇总。")],
    )
    endpoint_block(
        doc,
        "7.4 代理批发商品列表",
        "GET",
        "/buyer/agent/goods/wholesale",
        "是",
        "代理业务商品查询（通常结合批发规则）",
        "代理-批发页商品列表。",
        [
            ["keyword", "Query / string", "否", "商品名搜索。"],
            ["categoryId", "Query / string", "否", "分类筛选。"],
            ["storeId", "Query / string", "否", "店铺筛选。"],
        ],
        [
            ["goodsId", "string", "商品 ID。"],
            ["wholesaleList", "array", "阶梯批发价规则。"],
            ["minWholesaleNum", "int", "起批量。"],
        ],
        [("和普通商品搜索页不同", "代理批发页更强调阶梯价、起批量、店铺关系。")],
    )
    endpoint_block(
        doc,
        "7.5 采购对账列表",
        "GET",
        "/buyer/agent/procurement/reconciliation",
        "是",
        "DB 对账单聚合",
        "代理商采购对账页。",
        [
            ["dateRange", "Query / string", "否", "对账周期。"],
            ["status", "Query / string", "否", "对账状态。"],
            ["keyword", "Query / string", "否", "订单号/店铺名/供货商名。"],
        ],
        [
            ["id", "string", "对账单 ID。"],
            ["purchaseAmount", "number", "采购金额。"],
            ["purchaseCount", "int", "采购单数。"],
            ["settlementStatus / reconciliationStatus", "enum", "对账/结算状态。"],
        ],
        [("状态以后端枚举为准", "不要前端自造“已结清/未结清”并覆盖原始状态。")],
    )
    endpoint_block(
        doc,
        "7.6 资金对账列表",
        "GET",
        "/buyer/agent/fund/reconciliation",
        "是",
        "DB 对账单聚合",
        "代理商资金对账页。",
        [
            ["serviceType", "Query / string", "否", "资金流水业务类型，例如充值、提现、余额支付、退款、佣金。"],
            ["startDate", "Query / string", "否", "开始日期，格式 yyyy-MM-dd。"],
            ["endDate", "Query / string", "否", "结束日期，格式 yyyy-MM-dd。"],
        ],
        [
            ["id", "string", "资金流水 ID。"],
            ["memberId", "string", "代理商会员 ID。"],
            ["memberName", "string", "代理商会员名称。"],
            ["money", "number", "变动金额。"],
            ["serviceType", "string", "业务类型。"],
            ["detail", "string", "业务描述。"],
            ["createTime", "datetime", "创建时间。"],
        ],
        [("首期不提供下载文件", "当前买家端只支持分页、汇总、详情，不提供对账单下载接口。")],
    )

    add_h1(doc, "8. 供货商页面")
    endpoint_block(
        doc,
        "8.1 商品列表",
        "GET",
        "/store/goods/goods/list",
        "是（SUPPLIER）",
        "DB 商品查询 + 商家域权限",
        "供货商商品管理列表。",
        [
            ["goodsName", "Query / string", "否", "商品名称。"],
            ["marketEnable", "Query / string", "否", "上下架状态。"],
            ["pageNumber", "Query / int", "否", "页码。"],
            ["pageSize", "Query / int", "否", "分页大小。"],
        ],
        [
            ["goodsId", "string", "商品 ID。"],
            ["goodsName", "string", "商品名称。"],
            ["marketEnable", "string", "上下架状态。"],
            ["quantity", "int", "库存。"],
            ["salesNum", "int", "销量。"],
        ],
        [("注意是商家域 token", "必须先走 /store/passport/login/selectIdentity 获取 STORE 域 token。")],
    )
    endpoint_block(
        doc,
        "8.2 商品详情（商家端）",
        "GET",
        "/store/goods/goods/get/{id}",
        "是（SUPPLIER）",
        "DB 商品详情",
        "供货商商品编辑页回显。",
        [["id", "Path / string", "是", "商品 ID。"]],
        [
            ["goodsName", "string", "商品名称。"],
            ["goodsGalleryList", "array", "商品图。"],
            ["skuList", "array", "SKU 列表。"],
            ["marketEnable", "string", "上下架状态。"],
            ["wholesaleList", "array", "阶梯批发价。"],
        ],
        [],
    )
    endpoint_block(
        doc,
        "8.3 创建商品",
        "POST",
        "/store/goods/goods/create",
        "是（SUPPLIER）",
        "DB 商品写入 + SKU + 批发规则",
        "供货端发布商品。",
        [
            ["goodsName", "Body / string", "是", "商品名称。"],
            ["categoryId", "Body / string", "是", "商品分类 ID。"],
            ["goodsGalleryList", "Body / array", "是", "商品图片列表。"],
            ["skuList", "Body / array", "是", "SKU 列表。"],
            ["wholesaleList", "Body / array", "否", "批发阶梯价列表。"],
            ["marketEnable", "Body / string", "是", "上下架状态。"],
        ],
        [["goodsId", "string", "新商品 ID。"]],
        [("批发商品核心在 wholesaleList", "前端要把起批量和对应价格完整带给后端。")],
    )
    endpoint_block(
        doc,
        "8.4 更新商品",
        "PUT",
        "/store/goods/goods/update/{goodsId}",
        "是（SUPPLIER）",
        "DB 商品更新 + 缓存失效 + 索引刷新",
        "编辑商品。",
        [
            ["goodsId", "Path / string", "是", "商品 ID。"],
            ["goodsOperationDTO", "Body / object", "是", "完整商品编辑对象。"],
        ],
        [],
        [("编辑后前台可能有短缓存", "前端如果立刻回商品详情页验证，请考虑缓存刷新延迟。")],
    )
    endpoint_block(
        doc,
        "8.5 订单列表（商家端）",
        "GET",
        "/store/order/order",
        "是（SUPPLIER）",
        "DB 订单查询 + 商家域权限",
        "供货商订单管理页。",
        [
            ["orderStatus", "Query / string", "否", "订单状态筛选。"],
            ["keyword", "Query / string", "否", "订单号/收货人/商品关键词。"],
        ],
        [
            ["orderSn", "string", "订单号。"],
            ["orderStatus", "enum", "订单状态。"],
            ["memberName", "string", "下单人。"],
            ["flowPrice", "number", "订单金额。"],
        ],
        [],
    )
    endpoint_block(
        doc,
        "8.6 商家发货",
        "PUT",
        "/store/order/order/{orderSn}/deliver",
        "是（SUPPLIER）",
        "DB 订单状态流转",
        "商家手动发货。",
        [
            ["orderSn", "Path / string", "是", "订单号。"],
            ["物流参数", "Body/Form / object", "视实现而定", "通常包括物流公司、单号、包裹信息。"],
        ],
        [],
        [("核销类订单不走普通发货", "到店核销订单要走核销流程，不要误调发货接口。")],
    )
    endpoint_block(
        doc,
        "8.7 待核销订单查码",
        "GET",
        "/store/order/order/getOrderByVerificationCode/{verificationCode}",
        "是（SUPPLIER）",
        "DB 订单查询",
        "通过核销码查询订单。",
        [["verificationCode", "Path / string", "是", "核销码。"]],
        [
            ["orderSn", "string", "订单号。"],
            ["verificationCode", "string", "核销码。"],
            ["orderStatus", "enum", "通常应为待核销状态。"],
        ],
        [("用于门店核销前校验", "前端扫码后先查单，再做核销动作。")],
    )
    endpoint_block(
        doc,
        "8.8 核销订单",
        "PUT",
        "/store/order/order/take/{orderSn}/{verificationCode}",
        "是（SUPPLIER）",
        "DB 订单状态流转 + 核销记录",
        "执行核销。",
        [
            ["orderSn", "Path / string", "是", "订单号。"],
            ["verificationCode", "Path / string", "是", "核销码。"],
        ],
        [],
        [("核销成功后订单完成", "前端成功页应按已核销/已完成状态刷新。")],
    )
    endpoint_block(
        doc,
        "8.9 数据概览",
        "GET",
        "/store/dashboard/overview",
        "是（SUPPLIER）",
        "DB 聚合统计",
        "供货商工作台/数据概览。",
        [],
        [
            ["storeIndexStatistics", "object", "店铺首页概览，包括商品数、订单总额、待处理事项、库存预警等。"],
            ["orderOverview", "object", "订单概览，包括下单、支付、退款、转化率等统计。"],
            ["goodsRankList", "array", "热卖商品排行。"],
        ],
        [("当前版本范围", "供货商工作台一期仅支持营业总览与商品排行，不包含销售团队、地区统计。")],
    )
    endpoint_block(
        doc,
        "8.10 资产概览",
        "GET",
        "/store/dashboard/assets",
        "是（SUPPLIER）",
        "DB 资产汇总",
        "供货商资产页。",
        [],
        [
            ["totalIncome", "number", "累计收益。"],
            ["settledAmount", "number", "已结算金额。"],
            ["withdrawableAmount", "number", "可提现金额。"],
            ["frozenAmount", "number", "冻结金额。"],
        ],
        [],
    )

    add_h1(doc, "9. 统一状态与前端传值建议")
    add_table(
        doc,
        [
            ["状态类字段", "常见值", "前端说明"],
            ["identityCode", "CONSUMER / AGENT / SUPPLIER", "登录身份码；不要自己翻译成数字。"],
            ["identityStatus", "AVAILABLE / NOT_OPENED / PENDING / REJECTED / DISABLED", "身份卡片状态。"],
            ["nextAction", "NONE / APPLY_AGENT / APPLY_STORE / RESUBMIT / CONTACT_ADMIN", "身份不可进入时的引导动作。"],
            ["salesModel", "RETAIL / WHOLESALE", "商品销售模式。"],
            ["orderStatus", "UNPAID / PAID / DELIVERED / COMPLETED / CANCELLED 等", "订单状态由后端统一返回。"],
            ["orderTag", "ALL / WAIT_PAY / WAIT_SHIP / WAIT_ROG / COMPLETE / CANCELLED", "订单列表 tab 常用标签。"],
            ["deliveryMethod", "SELF_PICK_UP / LOCAL_TOWN_DELIVERY / LOGISTICS / VIRTUAL", "配送方式。"],
            ["afterSaleType", "RETURN_MONEY / RETURN_GOODS", "售后类型，当前没有换货枚举。"],
            ["refundWay", "ORIGINAL / OFFLINE", "退款方式。"],
            ["promotionStatus", "NEW / START / END / CLOSE", "活动状态。"],
            ["memberCouponStatus", "NEW / USED / EXPIRE / CLOSED", "会员券状态。"],
            ["storeAuditStatus", "DRAFT / SUBMITTED / APPROVED / REJECTED / FROZEN", "店铺审核状态。"],
            ["messageStatus", "UN_READY / ALREADY_READY / ALREADY_REMOVE", "消息未读/已读/已删除。"],
        ],
        [2100, 2800, 4460],
    )

    add_h1(doc, "10. 联调易错点清单")
    add_bullets(
        doc,
        [
            ("搜索页误用 DB 接口", "商品搜索和筛选不要用 /buyer/goods/goods，应该用 /buyer/goods/goods/es。"),
            ("把 loginSessionToken 当 accessToken", "登录认证成功后还要选身份，最终 Bearer token 从 selectIdentity 返回。"),
            ("优惠券 ID 混淆", "领取券通常用 couponId，结算选择券通常用 memberCouponId。"),
            ("地址字段假设错误", "项目地址常用地区路径字段，不一定只有 provinceId/cityId/areaId。"),
            ("售后路径参数误解", "售后申请里 orderItemSn 和 orderSn 不是一回事。"),
            ("历史动作接口用 GET", "例如部分领券、选择动作仍是 GET，前端不要主观改成 POST。"),
            ("商家域和买家域 token 混用", "SUPPLIER 要用 STORE 域 token，CONSUMER/AGENT 要用 MEMBER 域 token。"),
            ("缓存更新误判", "商品、SKU、购物车、热词、首页等都存在缓存或聚合层，看到短暂延迟先排查缓存刷新。"),
        ],
    )

    add_h1(doc, "11. 结论")
    add_paragraph(
        doc,
        "本项目已经明确区分了“重搜索走 ES、重交互走缓存上下文、最终业务数据落数据库”的链路。前端联调时最重要的是按页面使用正确接口，不要图省事把数据库列表接口替代 ES 搜索接口，也不要本地重算订单、优惠、运费和售后状态。只要按本文档的页面目录逐个联调，基本可以把绝大多数反复沟通的问题提前消化掉。",
        after=8,
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"SAVED::{OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()
